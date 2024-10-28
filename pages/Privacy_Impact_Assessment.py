import streamlit as st
from io import BytesIO
from fpdf import FPDF
from docx import Document

# Function to generate the Privacy Impact Assessment content
def generate_pia_content(answers):
    content = f"""
    Privacy Impact Assessment (PIA) Report
    
    1. What is the purpose of the project? 
    Answer: {answers['purpose']}

    2. What types of personal data will be collected? 
    Answer: {answers['data_types']}

    3. Who will have access to the data? 
    Answer: {answers['access']}

    4. How will the data be protected? 
    Answer: {answers['protection']}

    5. How long will the data be retained? 
    Answer: {answers['retention']}

    6. Are there any risks to data privacy, and how will they be mitigated? 
    Answer: {answers['risks']}

    7. Additional Comments: 
    Answer: {answers['comments']}
    """
    return content

# Function to generate a PDF version of the report
def generate_pdf(content):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, content)

    # Output the PDF as a byte stream
    pdf_output = BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin1')
    pdf_output.write(pdf_data)
    pdf_output.seek(0)
    return pdf_output

# Function to generate a Word (docx) version of the report
def generate_word(content):
    doc = Document()
    doc.add_heading('Privacy Impact Assessment (PIA) Report', 0)

    # Split the content into lines and add each line as a paragraph
    for line in content.split("\n"):
        doc.add_paragraph(line)

    # Save the document to a BytesIO object
    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

# Streamlit app
def app():
    st.title("Privacy Impact Assessment (PIA)")

    st.write("Please fill out the following questions to complete the Privacy Impact Assessment.")

    # Collect user input
    purpose = st.text_area("1. What is the purpose of the project?", "")
    data_types = st.text_area("2. What types of personal data will be collected?", "")
    access = st.text_area("3. Who will have access to the data?", "")
    protection = st.text_area("4. How will the data be protected?", "")
    retention = st.text_area("5. How long will the data be retained?", "")
    risks = st.text_area("6. Are there any risks to data privacy, and how will they be mitigated?", "")
    comments = st.text_area("7. Additional Comments:", "")

    # Gather the inputs into a dictionary
    answers = {
        "purpose": purpose,
        "data_types": data_types,
        "access": access,
        "protection": protection,
        "retention": retention,
        "risks": risks,
        "comments": comments,
    }

    if st.button("Generate Report"):
        # Generate the PIA content
        content = generate_pia_content(answers)

        # Display the generated content
        st.subheader("Generated Privacy Impact Assessment Report")
        st.text_area("PIA Report", value=content, height=400)

        # Provide options to download the report
        txt_data = content.encode("utf-8")

        st.download_button(
            label="Download as TXT",
            data=txt_data,
            file_name="PIA_Report.txt",
            mime="text/plain"
        )

        # Generate and provide the PDF
        pdf = generate_pdf(content)
        st.download_button(
            label="Download as PDF",
            data=pdf,
            file_name="PIA_Report.pdf",
            mime="application/pdf"
        )

        # Generate and provide the Word document
        word_doc = generate_word(content)
        st.download_button(
            label="Download as Word",
            data=word_doc,
            file_name="PIA_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    app()