import streamlit as st
from io import BytesIO
from fpdf import FPDF
from docx import Document
import datetime

# Function to generate the Data Breach Response Plan content
def generate_response_plan(details):
    content = f"""
    Data Breach Response Plan

    1. Incident Overview:
    Incident Type: {details['incident_type']}
    Date of Breach: {details['breach_date']}
    Affected Systems: {details['affected_systems']}
    Description: {details['description']}

    2. Data Affected:
    Types of Data Affected: {details['data_affected']}

    3. Detection and Response:
    How was the breach detected: {details['detection_method']}
    Immediate Response: {details['immediate_response']}

    4. Mitigation Steps:
    Mitigation Actions: {details['mitigation_steps']}

    5. Communication:
    Internal Communication Plan: {details['internal_communication']}
    External Communication Plan (e.g., authorities, customers): {details['external_communication']}

    6. Preventive Measures:
    Steps to prevent future breaches: {details['preventive_measures']}
    
    7. Assigned Responsibility:
    Assigned Team/Person: {details['responsible_person']}

    8. Review and Improvement:
    Future Improvements and Lessons Learned: {details['improvements']}
    
    9. Conclusion:
    Summary of Actions Taken: {details['summary']}
    """
    return content

# Function to generate a PDF version of the Data Breach Response Plan
def generate_pdf(content):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, content)

    pdf_output = BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin1')
    pdf_output.write(pdf_data)
    pdf_output.seek(0)
    return pdf_output

# Function to generate a Word (docx) version of the Data Breach Response Plan
def generate_word(content):
    doc = Document()
    doc.add_heading('Data Breach Response Plan', 0)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

# Streamlit app
def app():
    st.title("Data Breach Response Plan Generator")

    st.write("Fill out the fields below to generate your Data Breach Response Plan.")

    # Collect user input
    incident_type = st.text_input("Incident Type", value="Unauthorized access")
    breach_date = st.date_input("Date of Breach", value=datetime.date.today())
    affected_systems = st.text_area("Affected Systems", value="Customer database, Financial records")
    description = st.text_area("Incident Description", value="An unauthorized party accessed customer data.")
    data_affected = st.text_area("Types of Data Affected", value="Personal information, Financial data")
    detection_method = st.text_area("How was the breach detected?", value="The breach was detected via routine security monitoring.")
    immediate_response = st.text_area("Immediate Response Actions", value="Disabled access, locked down affected systems, and started internal investigation.")
    mitigation_steps = st.text_area("Mitigation Actions", value="Notified affected individuals, initiated password resets, implemented stronger authentication protocols.")
    internal_communication = st.text_area("Internal Communication Plan", value="Notify senior leadership, legal team, and IT team.")
    external_communication = st.text_area("External Communication Plan", value="Notify affected customers and report the incident to relevant authorities.")
    preventive_measures = st.text_area("Preventive Measures", value="Conduct regular security audits, improve network monitoring, and enhance staff training.")
    responsible_person = st.text_input("Assigned Responsibility (Team or Person)", value="John Doe, IT Security Team")
    improvements = st.text_area("Improvements and Lessons Learned", value="Review and enhance access control policies, conduct more frequent security drills.")
    summary = st.text_area("Summary of Actions Taken", value="All affected systems were secured, customers notified, and preventive steps were taken.")

    # Collect input into a dictionary
    details = {
        "incident_type": incident_type,
        "breach_date": breach_date,
        "affected_systems": affected_systems,
        "description": description,
        "data_affected": data_affected,
        "detection_method": detection_method,
        "immediate_response": immediate_response,
        "mitigation_steps": mitigation_steps,
        "internal_communication": internal_communication,
        "external_communication": external_communication,
        "preventive_measures": preventive_measures,
        "responsible_person": responsible_person,
        "improvements": improvements,
        "summary": summary
    }

    if st.button("Generate Response Plan"):
        # Generate the data breach response plan content
        content = generate_response_plan(details)

        # Display the data breach response plan content
        st.subheader("Generated Data Breach Response Plan")
        st.text_area("Response Plan", value=content, height=400)

        # Download options for TXT
        st.download_button(
            label="Download as TXT",
            data=content.encode("utf-8"),  # Encode as utf-8 string
            file_name=f"Data_Breach_Response_Plan_{incident_type}.txt",
            mime="text/plain"
        )

        # Download options for PDF
        pdf = generate_pdf(content)
        st.download_button(
            label="Download as PDF",
            data=pdf,
            file_name=f"Data_Breach_Response_Plan_{incident_type}.pdf",
            mime="application/pdf"
        )

        # Download options for Word
        word_doc = generate_word(content)
        st.download_button(
            label="Download as Word",
            data=word_doc,
            file_name=f"Data_Breach_Response_Plan_{incident_type}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    app()