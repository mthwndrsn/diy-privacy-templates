import streamlit as st
from io import BytesIO
from fpdf import FPDF
from docx import Document
from datetime import datetime

# Function to generate the full privacy policy text based on inputs
def generate_privacy_policy(company_name, website_url, contact_email, data_types, collection_methods, usage_purposes,
                            third_party_sharing, third_party_names, security_measures, user_rights, retention_period, policy_changes):
    
    policy = f"""
    Privacy Policy for {company_name}
    
    At {company_name}, accessible from {website_url}, one of our main priorities is the privacy of our visitors. 
    This Privacy Policy document contains types of information that is collected and recorded by {company_name} 
    and how we use it.

    If you have additional questions or require more information about our Privacy Policy, do not hesitate to 
    contact us through email at {contact_email}.
    
    ## 1. Information We Collect
    
    We collect the following types of information:
    {', '.join(data_types)}.

    We collect this information through the following methods:
    {', '.join(collection_methods)}.

    ## 2. How We Use Your Information
    
    The information we collect is used for various purposes, including:
    {', '.join(usage_purposes)}.

    ## 3. Sharing Your Information
    
    {f"We share your information with third-party services, including: {', '.join(third_party_names)}." if third_party_sharing == 'Yes' else "We do not share your personal data with third-party services without your consent."}

    ## 4. Data Security Measures
    
    We implement the following data security measures to protect your information:
    {security_measures}.
    
    ## 5. Your Data Protection Rights
    
    You have the right to request copies of your personal data, rectify any inaccurate information, and request the erasure of your personal data under certain conditions. The rights you are entitled to include:
    {user_rights}.
    
    ## 6. Data Retention
    
    We retain personal data for {retention_period}. After this period, we delete or anonymize the data.

    ## 7. Changes to Our Privacy Policy
    
    {policy_changes}

    For more detailed information, feel free to contact us at {contact_email}.
    """

    return policy

# Function to generate PDF version of the privacy policy
def generate_pdf(privacy_policy):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, privacy_policy)
    
    # Output the PDF as a byte string
    pdf_output = BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin1')  # 'S' to return as string
    pdf_output.write(pdf_data)
    pdf_output.seek(0)  # Move to the beginning of the file-like object
    return pdf_output

# Function to generate Word (docx) version of the privacy policy with formatting
def generate_word(privacy_policy):
    doc = Document()
    doc.add_heading('Privacy Policy', 0)

    # Split the policy text into lines and process each line
    for line in privacy_policy.split("\n"):
        # Check for lines with markdown-style bold (**text**)
        if "**" in line:
            # Split by '**' to identify bold sections
            parts = line.split("**")
            paragraph = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Regular text
                    paragraph.add_run(part)
                else:
                    # Bold text
                    paragraph.add_run(part).bold = True
        else:
            doc.add_paragraph(line)

    # Save the document to a BytesIO object
    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)  # Move to the beginning of the file-like object
    return word_output

# Streamlit app
def app():
    st.title("Robust Privacy Policy Generator")

    st.write("Generate a custom, comprehensive privacy policy for your business by filling out the details below.")

    # User input fields for the expanded privacy policy
    company_name = st.text_input("Company Name", value="My Company")
    website_url = st.text_input("Website URL", value="https://www.mycompany.com")
    contact_email = st.text_input("Contact Email", value="info@mycompany.com")

    st.header("1. Information We Collect")
    data_types = st.multiselect("What types of data do you collect?", 
                                ['Personal Information', 'Email Address', 'Payment Information', 'Usage Data', 
                                 'Cookies', 'IP Address'], default=['Personal Information', 'Email Address'])

    collection_methods = st.multiselect("How do you collect this information?", 
                                        ['Website Forms', 'Cookies', 'Analytics Tools', 'Direct Interactions', 'Third Parties'], 
                                        default=['Website Forms', 'Cookies'])

    st.header("2. How We Use Your Information")
    usage_purposes = st.multiselect("What do you use this data for?", 
                                    ['Personalizing User Experience', 'Processing Payments', 'Improving Website Functionality', 
                                     'Customer Support', 'Marketing Purposes'], default=['Personalizing User Experience', 'Customer Support'])

    st.header("3. Sharing Information with Third Parties")
    third_party_sharing = st.radio("Do you share data with third-party services?", ["Yes", "No"])
    third_party_names = []
    if third_party_sharing == "Yes":
        third_party_names = st.text_area("List the third-party services you share data with (e.g., Google Analytics, Stripe):").split(',')

    st.header("4. Data Security Measures")
    security_measures = st.text_area("What data security measures do you have in place?", 
                                     "Encryption, Secure Data Storage, Regular Audits")

    st.header("5. User Data Protection Rights")
    user_rights = st.text_area("What are the rights of the users with respect to their data?", 
                               "Right to Access, Right to Rectification, Right to Erasure, Right to Data Portability")

    st.header("6. Data Retention")
    retention_period = st.text_input("How long do you retain personal data?", "1 year")

    st.header("7. Changes to Our Privacy Policy")
    policy_changes = st.text_area("How will you notify users of changes to the privacy policy?", 
                                  "We will notify users via email and update the policy on our website.")

    # Get the current date
    current_date = datetime.now().strftime("%Y-%m-%d")

    if st.button("Generate Privacy Policy"):
        # Generate the privacy policy
        privacy_policy = generate_privacy_policy(company_name, website_url, contact_email, data_types, collection_methods, 
                                                 usage_purposes, third_party_sharing, third_party_names, security_measures, 
                                                 user_rights, retention_period, policy_changes)
        
        # Display the generated policy
        st.subheader("Generated Privacy Policy")
        st.text_area("Privacy Policy", value=privacy_policy, height=400)

        # Option to download as .txt file
        st.download_button(
            label="Download as .txt",
            data=privacy_policy,
            file_name=f"{company_name}_privacy_policy_{current_date}.txt",
            mime="text/plain",
        )

        # Option to download as PDF
        pdf = generate_pdf(privacy_policy)
        st.download_button(
            label="Download as PDF",
            data=pdf,
            file_name=f"{company_name}_privacy_policy_{current_date}.pdf",
            mime="application/octet-stream"
        )

        # Option to download as Word (.docx)
        word_doc = generate_word(privacy_policy)
        st.download_button(
            label="Download as Word",
            data=word_doc,
            file_name=f"{company_name}_privacy_policy_{current_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    app()