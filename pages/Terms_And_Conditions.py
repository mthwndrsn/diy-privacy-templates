import streamlit as st
from io import BytesIO
from fpdf import FPDF
from docx import Document

# Function to generate the Terms and Conditions content based on user input
def generate_terms_content(details):
    content = f"""
    Terms and Conditions

    Effective Date: {details['effective_date']}
    
    1. Introduction
    Welcome to {details['company_name']}! These terms and conditions outline the rules and regulations for the use of {details['company_name']}'s Website, located at {details['website_url']}.
    
    2. Intellectual Property Rights
    Unless otherwise stated, {details['company_name']} and/or its licensors own the intellectual property rights for all material on {details['website_url']}. All intellectual property rights are reserved. You may access this from {details['website_url']} for your own personal use subjected to restrictions set in these terms and conditions.
    
    3. Restrictions
    You are specifically restricted from all of the following:
    - Publishing any Website material in any other media;
    - Selling, sublicensing, and/or otherwise commercializing any Website material;
    - Publicly performing and/or showing any Website material;
    - Using this Website in any way that is or may be damaging to this Website;
    - Using this Website in any way that impacts user access to this Website;
    - Engaging in any data mining, data harvesting, data extracting, or any other similar activity in relation to this Website;
    - Using this Website to engage in any advertising or marketing.

    4. User Content
    In these Website Standard Terms and Conditions, "User Content" shall mean any audio, video text, images, or other material you choose to display on this Website. By displaying it, you grant {details['company_name']} a non-exclusive, worldwide irrevocable, sub-licensable license to use, reproduce, adapt, publish, translate, and distribute it in any media.

    5. No warranties
    This Website is provided "as is," with all faults, and {details['company_name']} express no representations or warranties of any kind related to this Website or the materials contained on this Website.

    6. Limitation of liability
    In no event shall {details['company_name']}, nor any of its officers, directors, or employees, be held liable for anything arising out of or in any way connected with your use of this Website whether such liability is under contract. {details['company_name']} shall not be held liable for any indirect, consequential, or special liability arising out of or in any way related to your use of this Website.

    7. Governing Law & Jurisdiction
    These Terms will be governed by and interpreted in accordance with the laws of {details['jurisdiction']}, and you submit to the non-exclusive jurisdiction of the state and federal courts located in {details['jurisdiction']} for the resolution of any disputes.

    """
    return content

# Function to generate a PDF version of the terms and conditions
def generate_pdf(content):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, content)

    pdf_output = BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin1')
    pdf_output.write(pdf_data)
    pdf_output.seek(0)
    return pdf_output

# Function to generate a Word (docx) version of the terms and conditions
def generate_word(content):
    doc = Document()
    doc.add_heading('Terms and Conditions', 0)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

# Streamlit app
def app():
    st.title("Terms and Conditions Generator")

    st.write("Fill out the fields below to generate your Terms and Conditions document.")

    # Collect user input
    company_name = st.text_input("Company Name", value="Acme Corp")
    website_url = st.text_input("Website URL", value="https://www.acme.com")
    effective_date = st.date_input("Effective Date")
    jurisdiction = st.text_input("Governing Law & Jurisdiction", value="New York, USA")

    # Collect input into a dictionary
    details = {
        "company_name": company_name,
        "website_url": website_url,
        "effective_date": effective_date,
        "jurisdiction": jurisdiction,
    }

    if st.button("Generate Terms and Conditions"):
        # Generate the terms and conditions content
        content = generate_terms_content(details)

        # Display the terms and conditions content
        st.subheader("Generated Terms and Conditions")
        st.text_area("Terms and Conditions", value=content, height=400)

        # Download options for TXT
        st.download_button(
            label="Download as TXT",
            data=content.encode("utf-8"),  # Encode as utf-8 string
            file_name=f"Terms_and_Conditions_{company_name}.txt",
            mime="text/plain"
        )

        # Download options for PDF
        pdf = generate_pdf(content)
        st.download_button(
            label="Download as PDF",
            data=pdf,
            file_name=f"Terms_and_Conditions_{company_name}.pdf",
            mime="application/pdf"
        )

        # Download options for Word
        word_doc = generate_word(content)
        st.download_button(
            label="Download as Word",
            data=word_doc,
            file_name=f"Terms_and_Conditions_{company_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    app()